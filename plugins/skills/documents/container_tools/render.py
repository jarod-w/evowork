#!/usr/bin/env python3
"""内容 JSON → docx / md。08 §5.2 的第一个技能。

模型给**内容与层级**，字号 / 行距 / 页眉页脚 / 目录 / 表格样式由模板决定。

## 两条与别的技能不同的点

1. **md 是一等输出，不是降级。** 很多场景（写进仓库的说明、发给同事的草稿）要的就是 md，
   而 md 走内置渲染，不需要办公扩展 —— 这让"没装扩展"不等于"什么都干不了"。
2. **docx 的目录是一个域（field），不是一段文字。** python-docx 写进去的目录在 Word 里
   第一次打开时是空的，要按 F9 更新域才会出现。这件事必须写在文档里告诉用户，
   否则他会以为目录生成失败了 —— 所以渲染器在目录下方留一行提示。

退出码与共用骨架一致（见 `plugins/skills/_shared/evowork_skill.py`）。
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

SKILL_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(SKILL_ROOT.parent / "_shared"))

from evowork_skill import (  # noqa: E402
    ensure_office_runtime,
    EXIT_INVALID_CONTENT,
    EXIT_RUNTIME_MISSING,
    fail,
    load_template,
    read_content,
    resolve_asset,
    runtime_missing_message,
    succeeded,
    validate_content,
)

BLOCKS = ("heading", "paragraph", "bullets", "ordered", "table", "image", "quote", "pagebreak")

TOC_HINT = "（目录是 Word 的域：第一次打开时按 F9 或右键「更新域」即可显示页码。）"


def validate(content: dict, base_dir: Path) -> None:
    validate_content(
        content,
        SKILL_ROOT / "schema" / "content.schema.json",
        discriminator="block",
        known_kinds=BLOCKS,
    )

    # schema 表达不了的两条跨字段约束
    for index, block in enumerate(content["blocks"]):
        if block["block"] == "table":
            width = len(block["header"])
            for row_index, row in enumerate(block["rows"]):
                if len(row) != width:
                    fail(
                        EXIT_INVALID_CONTENT,
                        f"blocks/{index}/rows/{row_index}: 这一行有 {len(row)} 格，"
                        f"表头有 {width} 列，两者必须一样长。",
                    )
        if block["block"] == "image":
            resolve_asset(block["path"], base_dir, hint="先用 charts 技能生成它，再引用。")

    levels = [b["level"] for b in content["blocks"] if b["block"] == "heading"]
    for previous, current in zip(levels, levels[1:]):
        if current > previous + 1:
            # 层级跳跃（h1 直接到 h3）会让目录看起来缺一层，且 Word 的导航窗格会错位
            fail(
                EXIT_INVALID_CONTENT,
                f"标题层级从 {previous} 跳到了 {current}。层级只能逐级下降"
                f"（h{previous} 之后是 h{previous + 1}），否则目录会缺一层。",
            )


# ─────────────────────────── markdown（基础包，无需扩展）───────────────────────────


def render_md(content: dict, out_path: Path) -> None:
    lines: list[str] = [f"# {content['title']}", ""]
    if content.get("subtitle"):
        lines += [f"> {content['subtitle']}", ""]
    meta = " · ".join(x for x in (content.get("author"), content.get("date")) if x)
    if meta:
        lines += [meta, ""]
    if content.get("toc"):
        # md 的目录由阅读器生成，写一行标记比伪造一份会过期的目录好
        lines += ["<!-- toc -->", ""]

    for block in content["blocks"]:
        kind = block["block"]
        if kind == "heading":
            lines += ["#" * (block["level"] + 1) + " " + block["text"], ""]
        elif kind == "paragraph":
            lines += [block["text"], ""]
        elif kind == "bullets":
            lines += [f"- {item}" for item in block["items"]] + [""]
        elif kind == "ordered":
            lines += [f"{i + 1}. {item}" for i, item in enumerate(block["items"])] + [""]
        elif kind == "quote":
            lines += [f"> {block['text']}"]
            if block.get("cite"):
                lines += [f">", f"> —— {block['cite']}"]
            lines += [""]
        elif kind == "table":
            header = block["header"]
            lines += ["| " + " | ".join(header) + " |"]
            lines += ["| " + " | ".join("---" for _ in header) + " |"]
            for row in block["rows"]:
                lines += ["| " + " | ".join("" if c is None else str(c) for c in row) + " |"]
            if block.get("caption"):
                lines += ["", f"*{block['caption']}*"]
            lines += [""]
        elif kind == "image":
            lines += [f"![{block.get('caption', '')}]({block['path']})", ""]
        elif kind == "pagebreak":
            # md 没有分页概念。用水平线表达"这里原本是分页"，而不是静默丢掉
            lines += ["---", ""]

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


# ─────────────────────────── docx（需要办公扩展）───────────────────────────


def render_docx(content: dict, out_path: Path, base_dir: Path) -> None:
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ModuleNotFoundError:
        fail(EXIT_RUNTIME_MISSING, runtime_missing_message("office", "生成 docx"))

    template = load_template(SKILL_ROOT, content.get("template", "report"))
    styles = template["styles"]
    fonts = template["font"]

    document = docx.Document()
    section = document.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(template["page"]["margin_cm"]))

    def style_run(run, spec: dict) -> None:
        run.font.size = Pt(spec["size_pt"])
        run.font.bold = spec.get("bold", False)
        run.font.italic = spec.get("italic", False)
        if spec.get("color"):
            run.font.color.rgb = RGBColor.from_string(spec["color"])
        run.font.name = fonts["latin"]
        # 中文字体必须写进 rPr 的 eastAsia，否则中文会回落到默认字体（可能是方框）
        run._element.rPr.rFonts.set(qn("w:eastAsia"), fonts["cjk"])

    def add_paragraph(text: str, spec: dict):
        paragraph = document.add_paragraph()
        if spec.get("align") == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(spec.get("space_before_pt", 0))
        fmt.space_after = Pt(spec.get("space_after_pt", 0))
        if spec.get("line_spacing"):
            fmt.line_spacing = spec["line_spacing"]
        if spec.get("indent_cm"):
            fmt.left_indent = Cm(spec["indent_cm"])
        style_run(paragraph.add_run(text), spec)
        return paragraph

    def add_toc_field() -> None:
        """插入 TOC 域。python-docx 没有现成 API，只能拼 XML。"""
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    # 封面
    add_paragraph(content["title"], styles["title"])
    if content.get("subtitle"):
        add_paragraph(content["subtitle"], styles["subtitle"])
    meta = " · ".join(x for x in (content.get("author"), content.get("date")) if x)
    if meta:
        add_paragraph(meta, styles["subtitle"])
    if template.get("cover_page"):
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    if content.get("toc"):
        add_paragraph("目录", styles["h1"])
        add_toc_field()
        # 不写这一行的话，用户看到空目录会以为生成失败了（见文件头）
        add_paragraph(TOC_HINT, styles["caption"])
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    for block in content["blocks"]:
        kind = block["block"]
        if kind == "heading":
            add_paragraph(block["text"], styles[f"h{block['level']}"])
        elif kind == "paragraph":
            add_paragraph(block["text"], styles["body"])
        elif kind == "quote":
            add_paragraph(block["text"], styles["quote"])
            if block.get("cite"):
                add_paragraph(f"—— {block['cite']}", styles["caption"])
        elif kind in ("bullets", "ordered"):
            style_name = "List Bullet" if kind == "bullets" else "List Number"
            for item in block["items"]:
                paragraph = document.add_paragraph(style=style_name)
                style_run(paragraph.add_run(item), styles["body"])
        elif kind == "table":
            spec = template["table"]
            table = document.add_table(rows=1, cols=len(block["header"]))
            table.style = spec["style"]
            for index, text in enumerate(block["header"]):
                cell = table.rows[0].cells[index]
                cell.text = ""
                style_run(
                    cell.paragraphs[0].add_run(text),
                    {**styles["body"], "size_pt": spec["font_pt"], "bold": spec["header_bold"]},
                )
                if spec.get("header_fill"):
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), spec["header_fill"])
                    cell._tc.get_or_add_tcPr().append(shading)
            for row in block["rows"]:
                cells = table.add_row().cells
                for index, value in enumerate(row):
                    cells[index].text = ""
                    style_run(
                        cells[index].paragraphs[0].add_run("" if value is None else str(value)),
                        {**styles["body"], "size_pt": spec["font_pt"]},
                    )
            if block.get("caption"):
                add_paragraph(block["caption"], styles["caption"])
        elif kind == "image":
            path = resolve_asset(block["path"], base_dir, hint="先用 charts 技能生成它。")
            document.add_picture(
                str(path), width=Cm(block.get("width_cm", template["image_default_width_cm"]))
            )
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if block.get("caption"):
                add_paragraph(block["caption"], styles["caption"])
        elif kind == "pagebreak":
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    header_spec = template["header_footer"]
    if header_spec.get("header_text"):
        header = section.header.paragraphs[0]
        header.text = header_spec["header_text"].replace("{title}", content["title"])
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header_spec.get("footer_page_number"):
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="内容 JSON → docx / md")
    parser.add_argument("--content", required=True)
    parser.add_argument("--out", required=True, help="输出路径，扩展名决定格式（.docx / .md）")
    parser.add_argument("--validate-only", action="store_true")
    args = parser.parse_args()

    # 缺办公扩展的模块时换到扩展的解释器重跑（见 evowork_skill.ensure_office_runtime）。
    # --validate-only 不需要它，但提前换掉更简单，也让两条路径的行为一致
    ensure_office_runtime(("docx",))

    content_path = Path(args.content).resolve()
    content = read_content(content_path)
    validate(content, content_path.parent)

    if args.validate_only:
        succeeded(blocks=len(content["blocks"]), validated=True)
        return

    out_path = Path(args.out).resolve()
    suffix = out_path.suffix.lower()
    if suffix == ".md":
        render_md(content, out_path)
        succeeded(out_path, blocks=len(content["blocks"]), format="md")
        return
    if suffix != ".docx":
        fail(EXIT_INVALID_CONTENT, f"输出格式只支持 .docx 与 .md，收到 {suffix or '(无扩展名)'}。")

    render_docx(content, out_path, content_path.parent)
    succeeded(out_path, blocks=len(content["blocks"]), format="docx")


if __name__ == "__main__":
    main()
